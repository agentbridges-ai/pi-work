#!/usr/bin/env python3
"""Generate the deterministic Word-family preview showcase source."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "demo-user-space" / "office"
TEMPLATE = OUT_DIR / "Example Title.docx"
OUTPUT = OUT_DIR / "ONLYOFFICE Word Preview Showcase.docx"
ICON = ROOT / "web" / "public" / "icons" / "piwork-512.png"

NAVY = "12304A"
BLUE = "2F75B5"
TEAL = "00A6A6"
GOLD = "F2B134"
CORAL = "E76F51"
PALE = "EAF3F8"
GRID = "D7E0E7"


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = NAVY) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Aptos"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction: str, display: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)


def add_bookmark(paragraph, name: str, bookmark_id: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text: str, url: str | None = None, anchor: str | None = None) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if url:
        rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), rel_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_note_reference(paragraph, kind: str, note_id: str) -> None:
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FootnoteReference" if kind == "footnote" else "EndnoteReference")
    r_pr.append(style)
    run.append(r_pr)
    ref = OxmlElement(f"w:{kind}Reference")
    ref.set(qn("w:id"), note_id)
    run.append(ref)
    paragraph._p.append(run)


def add_list_paragraph(document: Document, text: str, num_id: int, level: int = 0):
    paragraph = document.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)
    return paragraph


def add_equation(paragraph) -> None:
    equation = parse_xml(
        f'<m:oMathPara {nsdecls("m", "w")}><m:oMath><m:r><m:t>E</m:t></m:r>'
        '<m:r><m:t>=</m:t></m:r><m:sSup><m:e><m:r><m:t>mc</m:t></m:r></m:e>'
        '<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup></m:oMath></m:oMathPara>'
    )
    paragraph._p.append(equation)


def add_content_controls(document: Document) -> None:
    body = document._element.body
    rich = parse_xml(
        f'<w:sdt {nsdecls("w")}><w:sdtPr><w:alias w:val="W14 Rich text"/>'
        '<w:tag w:val="W14-rich"/><w:id w:val="1401"/></w:sdtPr><w:sdtContent>'
        '<w:p><w:r><w:rPr><w:b/><w:color w:val="2F75B5"/></w:rPr>'
        '<w:t>W14 rich-text content control</w:t></w:r></w:p></w:sdtContent></w:sdt>'
    )
    checkbox = parse_xml(
        f'<w:sdt {nsdecls("w", "w14")}><w:sdtPr><w:alias w:val="W14 Checkbox"/>'
        '<w:tag w:val="W14-checkbox"/><w:id w:val="1402"/><w14:checkbox>'
        '<w14:checked w14:val="1"/><w14:checkedState w14:val="2612" w14:font="MS Gothic"/>'
        '<w14:uncheckedState w14:val="2610" w14:font="MS Gothic"/></w14:checkbox>'
        '</w:sdtPr><w:sdtContent><w:p><w:r><w:t>☒ W14 checked content control</w:t>'
        '</w:r></w:p></w:sdtContent></w:sdt>'
    )
    body.insert(len(body) - 1, rich)
    body.insert(len(body) - 1, checkbox)


def add_textbox(paragraph) -> None:
    pict = parse_xml(
        f'<w:r {nsdecls("w")} xmlns:v="urn:schemas-microsoft-com:vml"><w:pict><v:roundrect style="width:300pt;height:72pt" '
        'arcsize="12%" fillcolor="#EAF3F8" strokecolor="#2F75B5">'
        '<v:textbox inset="10pt,8pt,10pt,8pt"><w:txbxContent><w:p>'
        '<w:r><w:rPr><w:b/><w:color w:val="12304A"/></w:rPr>'
        '<w:t>W10 TEXT BOX · selectable DrawingML/VML fallback</w:t></w:r>'
        '</w:p></w:txbxContent></v:textbox></v:roundrect></w:pict></w:r>'
    )
    paragraph._p.append(pict)


def add_svg(document: Document, paragraph) -> None:
    svg = (
        '<svg xmlns="http://www.w3.org/2000/svg" width="640" height="180" viewBox="0 0 640 180">'
        '<rect width="640" height="180" rx="20" fill="#12304A"/>'
        '<circle cx="90" cy="90" r="46" fill="#00A6A6"/>'
        '<path d="M160 90h300" stroke="#F2B134" stroke-width="14" stroke-linecap="round"/>'
        '<text x="160" y="70" font-family="Arial" font-size="28" fill="white">W08 SVG vector resource</text>'
        '<text x="160" y="112" font-family="Arial" font-size="20" fill="#D7E0E7">non-zero intrinsic dimensions</text></svg>'
    ).encode()
    part = Part(PackURI("/word/media/showcase-vector.svg"), "image/svg+xml", svg, document.part.package)
    rel_id = document.part.relate_to(part, RT.IMAGE)
    drawing = parse_xml(
        f'<w:r {nsdecls("w", "wp", "a", "pic", "r")}><w:drawing><wp:inline distT="0" distB="0" distL="0" distR="0">'
        '<wp:extent cx="5486400" cy="1543050"/><wp:docPr id="808" name="W08 SVG" descr="W08 vector preview"/>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr><pic:cNvPr id="808" name="showcase-vector.svg" descr="W08 vector preview"/>'
        '<pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="' + rel_id + '"/><a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="5486400" cy="1543050"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
        '</a:graphicData></a:graphic></wp:inline></w:drawing></w:r>'
    )
    paragraph._p.append(drawing)


def make_last_picture_floating(paragraph) -> None:
    inline = paragraph._p.xpath(".//wp:inline")[-1]
    inline.tag = qn("wp:anchor")
    for key, value in {
        "simplePos": "0",
        "relativeHeight": "251659264",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
        "distT": "45720",
        "distB": "45720",
        "distL": "91440",
        "distR": "91440",
    }.items():
        inline.set(key, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "column")
    horizontal.append(parse_xml(f'<wp:posOffset {nsdecls("wp")}>0</wp:posOffset>'))
    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "paragraph")
    vertical.append(parse_xml(f'<wp:posOffset {nsdecls("wp")}>0</wp:posOffset>'))
    wrap = OxmlElement("wp:wrapSquare")
    wrap.set("wrapText", "bothSides")
    inline.insert(0, simple)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)
    doc_pr_index = next(i for i, child in enumerate(inline) if child.tag == qn("wp:docPr"))
    inline.insert(doc_pr_index, wrap)


def patch_note_parts(path: Path) -> None:
    replacements = {
        "word/footnotes.xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:footnotes {nsdecls("w")}><w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            '<w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t>W16 footnote: browser conversion must preserve note text.</w:t></w:r></w:p></w:footnote></w:footnotes>'
        ).encode(),
        "word/endnotes.xml": (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            f'<w:endnotes {nsdecls("w")}><w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
            '<w:endnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>'
            '<w:endnote w:id="1"><w:p><w:r><w:endnoteRef/></w:r><w:r><w:t>W16 endnote: semantic note resource retained.</w:t></w:r></w:p></w:endnote></w:endnotes>'
        ).encode(),
    }
    with ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    members.update(replacements)
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(tmp, "w", ZIP_DEFLATED) as target:
        for name, payload in members.items():
            target.writestr(name, payload)
    tmp.replace(path)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 28, NAVY), ("Subtitle", 12, "607785")):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    for name, size, color in (("Showcase Heading 1", 20, BLUE), ("Showcase Heading 2", 14, TEAL)):
        style = next((item for item in styles if item.name == name), None)
        if style is None:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    if "Marker" not in styles:
        marker = styles.add_style("Marker", WD_STYLE_TYPE.PARAGRAPH)
        marker.font.name = "Aptos"
        marker.font.size = Pt(8)
        marker.font.color.rgb = RGBColor.from_string(CORAL)


def clear_story(story) -> None:
    for child in list(story._element):
        story._element.remove(child)


def generate() -> None:
    document = Document(TEMPLATE)
    chart_paragraph = None
    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//c:chart"):
            chart_paragraph = deepcopy(paragraph._p)
            break
    if chart_paragraph is None:
        raise RuntimeError("Template native chart paragraph was not found")

    body = document._element.body
    section_properties = deepcopy(body.sectPr)
    for child in list(body):
        body.remove(child)
    body.append(section_properties)
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    header = section.header
    header.is_linked_to_previous = False
    clear_story(header)
    header_p = header.add_paragraph("ONLYOFFICE · WORD PREVIEW SHOWCASE · W01–W21")
    header_p.style = document.styles["Marker"]
    footer = section.footer
    footer.is_linked_to_previous = False
    clear_story(footer)
    footer_p = footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("派活 deterministic fixture · ")
    add_field(footer_p, "PAGE", "1")
    footer_p.add_run(" / ")
    add_field(footer_p, "NUMPAGES", "5")

    title = document.add_paragraph(style="Title")
    title.add_run("Word Preview Showcase")
    add_bookmark(title, "showcase_top", "42")
    document.add_paragraph(
        "A deterministic document-model corpus for typography, layout, drawings, charts, fields, notes and controls.",
        style="Subtitle",
    )
    if ICON.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(ICON), width=Inches(0.72))
    overview = document.add_table(rows=2, cols=3)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview.autofit = False
    for index, (label, value) in enumerate((("Families", "DOCX · DOC · ODT · RTF"), ("Acceptance", "package + editor model"), ("Build", "2026-07-22"))):
        set_cell_text(overview.cell(0, index), label, bold=True, color="FFFFFF")
        shade(overview.cell(0, index), NAVY)
        set_cell_text(overview.cell(1, index), value)
        shade(overview.cell(1, index), PALE)
    document.add_paragraph("W01 COVER · title/subtitle/header/footer/table styles", style="Marker")
    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Typography · Tables & media · Chart & semantics · Sections")

    document.add_page_break()
    document.add_paragraph("Typography and paragraph semantics", style="Showcase Heading 1")
    p = document.add_paragraph()
    p.add_run("W02 bold").bold = True
    p.add_run(" · italic").italic = True
    p.add_run(" · underline").underline = True
    p.add_run(" · strike").font.strike = True
    colored = p.add_run(" · coral color")
    colored.font.color.rgb = RGBColor.from_string(CORAL)
    highlighted = p.add_run(" · highlighted")
    highlighted.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = document.add_paragraph()
    p.add_run("W03 H₂O · x² · Ω ∑ ✓ · ")
    cjk = p.add_run("简体中文 · 日本語")
    cjk.font.name = "PingFang SC"
    cjk._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
    p.add_run(" · العربية · עברית").font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph("W04 paragraph border, shading, alignment, first-line indent and spacing are retained.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "12"), ("space", "4"), ("color", BLUE)):
        bottom.set(qn(f"w:{key}"), value)
    borders.append(bottom)
    p_pr.append(borders)
    document.add_paragraph("Lists", style="Showcase Heading 2")
    for text in ("W05 first bullet", "W05 second bullet", "W05 third bullet"):
        add_list_paragraph(document, text, 3)
    for text in ("W05 numbered step", "W05 numbered verification", "W05 numbered result"):
        add_list_paragraph(document, text, 2)
    nested = add_list_paragraph(document, "W05 nested level", 2, 1)
    nested.paragraph_format.left_indent = Inches(0.5)
    comment_run = document.add_paragraph().add_run("W15 comment anchor: semantic range survives conversion.")
    document.add_comment(comment_run, "W15 comment body retained by the editor model.", author="Piwork QA", initials="NQ")
    link_p = document.add_paragraph()
    add_hyperlink(link_p, "W13 authoritative DocumentServer", "https://github.com/ONLYOFFICE/DocumentServer")
    link_p.add_run(" · ")
    add_hyperlink(link_p, "return to cover", anchor="showcase_top")

    document.add_page_break()
    document.add_paragraph("Tables, images and drawings", style="Showcase Heading 1")
    table = document.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Marker", "Element", "Expected model", "Status")
    for index, label in enumerate(headers):
        set_cell_text(table.cell(0, index), label, bold=True, color="FFFFFF")
        shade(table.cell(0, index), NAVY)
    rows = (
        ("W06", "Merged cell", "gridSpan + fill", "READY"),
        ("W07", "Raster image", "decoded > 0×0", "READY"),
        ("W08", "SVG image", "image/svg+xml", "READY"),
        ("W10", "Text box", "drawing object", "READY"),
    )
    for row, values in enumerate(rows, 1):
        for col, value in enumerate(values):
            set_cell_text(table.cell(row, col), value, bold=(col == 0), color=TEAL if col == 0 else NAVY)
            shade(table.cell(row, col), "F7FAFC" if row % 2 else PALE)
    merged = table.cell(4, 1).merge(table.cell(4, 2))
    set_cell_text(merged, "W06 merged across two columns", bold=True)
    if ICON.exists():
        p = document.add_paragraph()
        run = p.add_run()
        run.add_picture(str(ICON), width=Inches(1.25))
        inline = p._p.xpath(".//wp:inline")[0]
        doc_pr = inline.find(qn("wp:docPr"))
        doc_pr.set("name", "W07 raster image")
        doc_pr.set("descr", "W07 raster image with alternative text")
        p.add_run("  W07 inline raster image · alt text · non-zero dimensions")
        floating = document.add_paragraph("W09 floating image with square wrapping. ")
        floating.add_run().add_picture(str(ICON), width=Inches(0.65))
        make_last_picture_floating(floating)
    add_svg(document, document.add_paragraph())
    add_textbox(document.add_paragraph())
    document.add_paragraph("W09 floating/wrapped image acceptance is asserted from anchor metadata in converted variants.", style="Marker")

    document.add_page_break()
    document.add_paragraph("Native chart and semantic objects", style="Showcase Heading 1")
    document.add_paragraph("W11 Native five-series chart · colored gradients must not collapse to black.", style="Marker")
    body.insert(len(body) - 1, chart_paragraph)
    caption = document.add_paragraph("Figure ")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(caption, "SEQ Figure \\* ARABIC", "1")
    caption.add_run(" — W11 conversion fidelity chart")
    for run in caption.runs:
        run.italic = True
    add_bookmark(caption, "_RefFigure1", "43")
    document.add_paragraph("Equation, fields and controls", style="Showcase Heading 2")
    eq = document.add_paragraph("W12 native equation: ")
    add_equation(eq)
    field_p = document.add_paragraph("W17/W19 fields: ")
    add_field(field_p, "DATE \\@ \"yyyy-MM-dd\"", "2026-07-22")
    field_p.add_run(" · cross-reference Figure ")
    add_field(field_p, "REF _RefFigure1 \\h", "1")
    note_p = document.add_paragraph("W16 note references: footnote")
    add_note_reference(note_p, "footnote", "1")
    note_p.add_run(" and endnote")
    add_note_reference(note_p, "endnote", "1")
    add_content_controls(document)

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    document.add_paragraph("Sections, columns and page geometry", style="Showcase Heading 1")
    document.add_paragraph("W18 landscape section · explicit section break · page-number/header/footer continuity.")
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    for index in range(1, 19):
        document.add_paragraph(
            f"W18 column paragraph {index}. This deterministic text validates column flow, paragraph spacing and section geometry."
        )
    document.add_paragraph("W20 legacy DOC OLE chart is covered by the focused Example Title.doc regression.", style="Marker")
    document.add_paragraph("W21 diagram coverage uses grouped shape/model fixtures in the presentation showcase.", style="Marker")

    props = document.core_properties
    props.title = "ONLYOFFICE Word Preview Showcase"
    props.subject = "Deterministic Word rendering and document-model fixture"
    props.author = "Piwork"
    props.keywords = "ONLYOFFICE, DOCX, preview, regression, document model"
    props.comments = "Generated by scripts/generate-office-word-showcase.py"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    patch_note_parts(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    generate()
